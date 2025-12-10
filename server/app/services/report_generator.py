import os
import base64
import io
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT
from reportlab.platypus import Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from PIL import Image, ImageDraw, ImageFont
import arabic_reshaper
from bidi.algorithm import get_display
import jdatetime
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import numpy as np

from app.core.config import settings
from app.services.file_manager import FileManager

class ReportGenerator:
    """Service for generating PDF and DOCX reports"""
    
    def __init__(self):
        self.file_manager = FileManager()
        self.fonts_dir = Path(settings.FONTS_DIR)
        
        # Register Persian font if available
        # Try Vazirmatn first (newer version), then fallback to Vazir
        self.has_persian_font = False
        font_candidates = [
            ("Vazirmatn-Regular.ttf", "Vazirmatn", "Vazirmatn-Bold.ttf"),
            ("Vazir-Regular.ttf", "Vazir", "Vazir-Bold.ttf"),
        ]
        
        for regular_file, font_name, bold_file in font_candidates:
            vazir_font = self.fonts_dir / regular_file
            vazir_bold = self.fonts_dir / bold_file
            
            if vazir_font.exists():
                try:
                    pdfmetrics.registerFont(TTFont(font_name, str(vazir_font)))
                    self.persian_font_name = font_name
                    self.has_persian_font = True
                    print(f"[REPORT] Persian font '{font_name}' loaded successfully from {vazir_font}")
                    
                    # Try to register bold variant if exists
                    if vazir_bold.exists():
                        try:
                            pdfmetrics.registerFont(TTFont(f"{font_name}-Bold", str(vazir_bold)))
                            self.persian_font_bold = f"{font_name}-Bold"
                            print(f"[REPORT] Persian bold font loaded from {vazir_bold}")
                        except Exception as e:
                            print(f"[REPORT] Warning: Could not load bold font: {e}")
                            self.persian_font_bold = font_name  # Fallback to regular
                    else:
                        self.persian_font_bold = font_name  # Fallback to regular
                    
                    break
                except Exception as e:
                    print(f"[REPORT] Warning: Could not load font '{font_name}': {e}")
        
        if not self.has_persian_font:
            print(f"[REPORT] Warning: No Persian font found in {self.fonts_dir}")
            print(f"[REPORT] Reports in Persian may not display correctly")
            self.persian_font_name = "Helvetica"
            self.persian_font_bold = "Helvetica-Bold"
    
    def _reshape_persian_text(self, text: str) -> str:
        """Reshape Persian/Arabic text for proper display"""
        try:
            reshaped = arabic_reshaper.reshape(text)
            bidi_text = get_display(reshaped)
            return bidi_text
        except Exception as e:
            print(f"[REPORT] Warning: Could not reshape text '{text}': {e}")
            return text
    
    def _draw_text_rtl(self, canvas_obj, text: str, x: float, y: float, max_width: float = None, is_rtl: bool = True):
        """Draw text with RTL support"""
        if is_rtl:
            text = self._reshape_persian_text(text)
            # For RTL, calculate position from right
            if max_width:
                text_width = canvas_obj.stringWidth(text, canvas_obj._fontname, canvas_obj._fontsize)
                x = x + max_width - text_width
        canvas_obj.drawString(x, y, text)
    
    def _draw_text_right_aligned(self, canvas_obj, text: str, x_right: float, y: float, is_rtl: bool = True):
        """Draw text aligned to the right edge"""
        if is_rtl:
            text = self._reshape_persian_text(text)
        text_width = canvas_obj.stringWidth(text, canvas_obj._fontname, canvas_obj._fontsize)
        canvas_obj.drawString(x_right - text_width, y, text)
    
    def _get_formatted_date(self, language: str = "en") -> str:
        """Get formatted date based on language (Jalali for Persian, Gregorian for English)"""
        now = datetime.utcnow()
        if language == 'fa':
            # Convert to Jalali (Persian) calendar
            j_date = jdatetime.datetime.fromgregorian(datetime=now)
            return j_date.strftime('%Y/%m/%d')
        else:
            # Gregorian calendar
            return now.strftime('%Y-%m-%d')
    
    def _create_temperature_histogram(self, markers: List[Dict[str, Any]], language: str = "en") -> io.BytesIO:
        """Create a temperature histogram from marker data"""
        if not markers:
            return None
        
        try:
            # Extract temperatures
            temps = [m.get('temperature', 0) for m in markers if m.get('temperature') is not None]
            
            if not temps:
                return None
            
            return self._create_histogram_from_temps(temps, language)
            
        except Exception as e:
            print(f"[REPORT] Error creating histogram from markers: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _create_temperature_histogram_from_csv(self, csv_data: str, language: str = "en") -> io.BytesIO:
        """Create a temperature histogram from CSV data"""
        try:
            import pandas as pd
            from io import StringIO
            
            # Parse CSV
            df = pd.read_csv(StringIO(csv_data))
            
            # Extract all temperature values (excluding headers)
            temps = df.values.flatten()
            temps = temps[~pd.isna(temps)]  # Remove NaN values
            
            if len(temps) == 0:
                return None
            
            return self._create_histogram_from_temps(temps, language)
            
        except Exception as e:
            print(f"[REPORT] Error creating histogram from CSV: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _create_histogram_from_temps(self, temps: List[float], language: str = "en") -> io.BytesIO:
        """Create a histogram from temperature array"""
        try:
            # Create figure
            fig, ax = plt.subplots(figsize=(6, 3))
            
            # Create histogram
            n_bins = min(30, max(10, len(temps) // 100))  # Dynamic bins
            ax.hist(temps, bins=n_bins, color='#ff6b6b', alpha=0.7, edgecolor='black')
            
            # For Persian, we need to set font that supports Persian
            if language == 'fa' and self.has_persian_font:
                # Use matplotlib's ability to render unicode
                # But we'll keep it simple with English labels for matplotlib
                # and add Persian title using PIL after
                ax.set_xlabel('Temperature (°C)', fontsize=10)
                ax.set_ylabel('Count', fontsize=10)
                ax.set_title('', fontsize=12)  # Will add Persian title separately
            else:
                ax.set_xlabel('Temperature (°C)', fontsize=10)
                ax.set_ylabel('Count', fontsize=10)
                ax.set_title('Temperature Distribution', fontsize=12, weight='bold')
            
            ax.grid(True, alpha=0.3)
            plt.tight_layout()
            
            # Save to BytesIO
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=100, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close(fig)
            
            # For Persian, add Persian text using PIL
            if language == 'fa':
                img = Image.open(img_buffer)
                draw = ImageDraw.Draw(img)
                
                # Try to load Persian font for PIL
                try:
                    font_path = self.fonts_dir / "Vazirmatn-Bold.ttf"
                    if font_path.exists():
                        pil_font = ImageFont.truetype(str(font_path), 16)
                    else:
                        pil_font = ImageFont.load_default()
                except:
                    pil_font = ImageFont.load_default()
                
                # Add Persian title
                title_text = "هیستوگرام دمایی"
                reshaped_title = self._reshape_persian_text(title_text)
                
                # Get text size and position
                bbox = draw.textbbox((0, 0), reshaped_title, font=pil_font)
                text_width = bbox[2] - bbox[0]
                text_x = (img.width - text_width) // 2
                text_y = 10
                
                # Draw text
                draw.text((text_x, text_y), reshaped_title, fill='black', font=pil_font)
                
                # Save modified image
                final_buffer = io.BytesIO()
                img.save(final_buffer, format='PNG')
                final_buffer.seek(0)
                return final_buffer
            
            return img_buffer
            
        except Exception as e:
            print(f"[REPORT] Error creating histogram: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _draw_table(self, canvas_obj, data: List[List[str]], x: float, y: float, col_widths: List[float], is_rtl: bool = False):
        """Draw a simple table"""
        if not data:
            return y
        
        # Table styling
        row_height = 20
        padding = 5
        
        # Draw rows
        current_y = y
        for row_idx, row in enumerate(data):
            current_x = x
            
            # Background for header
            if row_idx == 0:
                canvas_obj.setFillColorRGB(0.9, 0.9, 0.9)
                canvas_obj.rect(x, current_y - row_height, sum(col_widths), row_height, fill=1, stroke=1)
                canvas_obj.setFillColorRGB(0, 0, 0)
            else:
                canvas_obj.rect(x, current_y - row_height, sum(col_widths), row_height, fill=0, stroke=1)
            
            # Draw cells
            for col_idx, cell in enumerate(row):
                if is_rtl:
                    # Right-aligned for RTL
                    text = self._reshape_persian_text(str(cell)) if cell else ""
                    text_width = canvas_obj.stringWidth(text, canvas_obj._fontname, canvas_obj._fontsize)
                    canvas_obj.drawString(current_x + col_widths[col_idx] - text_width - padding, current_y - row_height + 5, text)
                else:
                    canvas_obj.drawString(current_x + padding, current_y - row_height + 5, str(cell))
                
                # Draw vertical line
                canvas_obj.line(current_x + col_widths[col_idx], current_y, current_x + col_widths[col_idx], current_y - row_height)
                current_x += col_widths[col_idx]
            
            current_y -= row_height
        
        return current_y
    
    def _base64_to_image(self, base64_string: str) -> Image.Image:
        """Convert base64 string to PIL Image"""
        if not base64_string:
            return None

        try:
            # Remove data URI prefix if present
            if ',' in base64_string:
                base64_string = base64_string.split(',')[1]

            image_data = base64.b64decode(base64_string)
            image = Image.open(io.BytesIO(image_data))
            return image
        except Exception as e:
            print(f"Error decoding base64 image: {e}")
            return None

    def _draw_professional_header(self, c, width: float, height: float, title: str, is_rtl: bool = False, page_num: int = 1):
        """Draw professional header with gradient background"""
        # Header background with gradient effect
        c.setFillColorRGB(0.1, 0.3, 0.6)  # Dark blue
        c.rect(0, height - 80, width, 80, fill=1, stroke=0)
        
        # Lighter blue stripe
        c.setFillColorRGB(0.2, 0.5, 0.8)
        c.rect(0, height - 85, width, 5, fill=1, stroke=0)
        
        # Title
        c.setFillColorRGB(1, 1, 1)  # White text
        font_bold = self.persian_font_bold if is_rtl and self.has_persian_font else 'Helvetica-Bold'
        c.setFont(font_bold, 20)
        
        if is_rtl:
            title = self._reshape_persian_text(title)
        c.drawCentredString(width/2, height - 45, title)
        
        # Subtitle
        c.setFont(font_bold if is_rtl else 'Helvetica', 10)
        subtitle = "گزارش تحلیل حرارتی" if is_rtl else "Thermal Analysis Report"
        if is_rtl:
            subtitle = self._reshape_persian_text(subtitle)
        c.drawCentredString(width/2, height - 65, subtitle)
        
        # Page number
        c.setFont('Helvetica', 8)
        c.drawRightString(width - 30, height - 20, f"Page {page_num}")
        
        # Reset color
        c.setFillColorRGB(0, 0, 0)
    
    def _draw_professional_footer(self, c, width: float, date_str: str, is_rtl: bool = False):
        """Draw professional footer"""
        # Footer line
        c.setStrokeColorRGB(0.2, 0.5, 0.8)
        c.setLineWidth(2)
        c.line(50, 50, width - 50, 50)
        
        # Date
        c.setFont('Helvetica', 9)
        c.setFillColorRGB(0.3, 0.3, 0.3)
        date_label = "تاریخ تولید گزارش: " if is_rtl else "Generated: "
        if is_rtl:
            date_label = self._reshape_persian_text(date_label)
        c.drawString(50, 35, f"{date_label}{date_str}")
        
        # Software name
        software = "Thermal Analyzer Pro"
        c.drawRightString(width - 50, 35, software)
        
        # Reset color
        c.setFillColorRGB(0, 0, 0)
        c.setStrokeColorRGB(0, 0, 0)
        c.setLineWidth(1)
    
    def _draw_section_header(self, c, title: str, y: float, width: float, is_rtl: bool = False):
        """Draw section header with background"""
        # Background
        c.setFillColorRGB(0.95, 0.95, 0.95)
        c.rect(50, y - 20, width - 100, 25, fill=1, stroke=1)
        
        # Title
        c.setFillColorRGB(0.1, 0.3, 0.6)
        font_bold = self.persian_font_bold if is_rtl and self.has_persian_font else 'Helvetica-Bold'
        c.setFont(font_bold, 12)
        
        if is_rtl:
            title = self._reshape_persian_text(title)
            c.drawRightString(width - 60, y - 15, title)
        else:
            c.drawString(60, y - 15, title)
        
        # Reset color
        c.setFillColorRGB(0, 0, 0)
        
        return y - 30

    def _draw_enhanced_table(self, c, data: List[List[str]], x: float, y: float, col_widths: List[float], is_rtl: bool = False):
        """Draw enhanced table with better styling"""
        if not data:
            return y
        
        row_height = 22
        padding = 5
        current_y = y
        
        # Draw rows
        for row_idx, row in enumerate(data):
            current_x = x
            
            # Header row styling
            if row_idx == 0:
                c.setFillColorRGB(0.2, 0.5, 0.8)  # Blue background
                c.rect(x, current_y - row_height, sum(col_widths), row_height, fill=1, stroke=0)
                c.setFillColorRGB(1, 1, 1)  # White text
                c.setFont(self.persian_font_bold if is_rtl and self.has_persian_font else 'Helvetica-Bold', 10)
            else:
                # Alternating row colors
                if row_idx % 2 == 0:
                    c.setFillColorRGB(0.98, 0.98, 0.98)
                else:
                    c.setFillColorRGB(1, 1, 1)
                c.rect(x, current_y - row_height, sum(col_widths), row_height, fill=1, stroke=0)
                c.setFillColorRGB(0, 0, 0)
                c.setFont(self.persian_font_name if is_rtl and self.has_persian_font else 'Helvetica', 9)
            
            # Draw border
            c.setStrokeColorRGB(0.7, 0.7, 0.7)
            c.rect(x, current_y - row_height, sum(col_widths), row_height, fill=0, stroke=1)
            
            # Draw cells
            for col_idx, cell in enumerate(row):
                cell_text = str(cell) if cell else ""
                
                if is_rtl:
                    cell_text = self._reshape_persian_text(cell_text)
                    text_width = c.stringWidth(cell_text, c._fontname, c._fontsize)
                    c.drawString(current_x + col_widths[col_idx] - text_width - padding, 
                               current_y - row_height + 6, cell_text)
                else:
                    c.drawString(current_x + padding, current_y - row_height + 6, cell_text)
                
                # Draw vertical separator
                if col_idx < len(col_widths) - 1:
                    c.setStrokeColorRGB(0.7, 0.7, 0.7)
                    c.line(current_x + col_widths[col_idx], current_y, 
                          current_x + col_widths[col_idx], current_y - row_height)
                
                current_x += col_widths[col_idx]
            
            current_y -= row_height
        
        # Reset
        c.setFillColorRGB(0, 0, 0)
        c.setStrokeColorRGB(0, 0, 0)
        
        return current_y

    def generate_pdf_report(
        self,
        project_id: str,
        metadata: Dict[str, Any],
        images: List[Dict[str, Any]],
        markers: List[Dict[str, Any]],
        regions: List[Dict[str, Any]],
        notes: Optional[str] = None,
        language: str = "en"
    ) -> str:
        """
        Generate PDF report with bilingual support (English/Persian)
        Returns: Path to generated PDF file
        """
        print(f"[REPORT] Starting PDF generation for project {project_id} in {language}")
        
        # Create report directory
        dirs = self.file_manager.create_project_directories(project_id)
        if not dirs:
            raise Exception("Failed to create project directories")
            
        report_dir = dirs["reports"]
        print(f"[REPORT] Report directory: {report_dir}")

        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        pdf_filename = f"report_{timestamp}.pdf"
        pdf_path = report_dir / pdf_filename

        # Create PDF
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4

        # Set font and text direction
        is_rtl = (language == 'fa')
        if language == 'fa' and self.has_persian_font:
            font_name = self.persian_font_name
            font_bold = self.persian_font_bold
        else:
            font_name = 'Helvetica'
            font_bold = 'Helvetica-Bold'
        
        print(f"[REPORT] Using font: {font_name}, RTL: {is_rtl}")

        # Professional Header
        page_num = 1
        self._draw_professional_header(c, width, height, title, is_rtl, page_num)

        # Project Information Section
        y = height - 100
        y = self._draw_section_header(c, 
            "اطلاعات پروژه" if is_rtl else "Project Information", 
            y, width, is_rtl)
        
        c.setFont(font_name, 10)

        # Labels based on language
        if is_rtl:
            info_labels = {
                "project_name": "پروژه",
                "operator": "اپراتور",
                "company": "شرکت",
                "customer": "مشتری",
                "camera_model": "مدل دوربین",
                "device": "دستگاه",
                "serial_no": "شماره سریال",
                "lens": "لنز",
                "measuring_site": "محل اندازه‌گیری",
                "task": "وظیفه",
                "date": "تاریخ",
                "time": "ساعت",
                "emissivity": "ضریب گسیل",
                "ambient_temp": "دمای محیط",
                "reflected_temp": "دمای بازتابی",
                "humidity": "رطوبت",
                "distance": "فاصله"
            }
        else:
            info_labels = {
                "project_name": "Project",
                "operator": "Operator",
                "company": "Company",
                "customer": "Customer",
                "camera_model": "Camera Model",
                "device": "Device",
                "serial_no": "Serial No.",
                "lens": "Lens",
                "measuring_site": "Measuring Site",
                "task": "Task",
                "date": "Date",
                "time": "Time",
                "emissivity": "Emissivity",
                "ambient_temp": "Ambient Temperature",
                "reflected_temp": "Reflected Temperature",
                "humidity": "Humidity",
                "distance": "Distance"
            }

        # Extract camera model from metadata or camera_model field
        camera_model = metadata.get("cameraModel", metadata.get("camera_model", metadata.get("device", "")))
        
        info_data = [
            ("project_name", metadata.get("project_name", "")),
            ("company", metadata.get("company", "")),
            ("customer", metadata.get("customer", "")),
            ("operator", metadata.get("operator", "")),
            ("camera_model", camera_model),
            ("device", metadata.get("device", metadata.get("device_model", ""))),
            ("serial_no", metadata.get("serial_no", metadata.get("serial_number", ""))),
            ("lens", metadata.get("lens", "")),
            ("measuring_site", metadata.get("measuring_site", metadata.get("location", ""))),
            ("date", metadata.get("date", self._get_formatted_date(language))),
            ("time", metadata.get("time", datetime.utcnow().strftime('%H:%M:%S'))),
        ]

        # Draw in two columns
        col1_x = 50
        col2_x = width / 2 + 20
        current_y = y
        col_switch = 0

        for key, value in info_data:
            if not value:  # Skip empty fields
                continue
                
            try:
                label = info_labels.get(key, key)
                if is_rtl:
                    # Right-aligned for RTL - label: value format
                    text = f"{label}: {value}"
                    reshaped_text = self._reshape_persian_text(text)
                    
                    if col_switch % 2 == 0:
                        self._draw_text_right_aligned(c, reshaped_text, width / 2 - 20, current_y, is_rtl=False)
                    else:
                        self._draw_text_right_aligned(c, reshaped_text, width - 50, current_y, is_rtl=False)
                        current_y -= 15
                else:
                    # Left-aligned for LTR
                    if col_switch % 2 == 0:
                        c.drawString(col1_x, current_y, f"{label}: {value}")
                    else:
                        c.drawString(col2_x, current_y, f"{label}: {value}")
                        current_y -= 15
                
                col_switch += 1
            except Exception as e:
                print(f"[REPORT] Warning: Could not render '{key}': {e}")
        
        # If odd number of items, move to next line
        if col_switch % 2 == 1:
            current_y -= 15
        
        y = current_y - 20
        
        # Add thermal parameters section if available
        thermal_params = [
            ("emissivity", metadata.get("emissivity")),
            ("ambient_temp", metadata.get("ambientTemp") or metadata.get("ambient_temp")),
            ("reflected_temp", metadata.get("reflectedTemp") or metadata.get("reflected_temp")),
            ("humidity", metadata.get("humidity")),
            ("distance", metadata.get("distance")),
        ]
        
        # Filter out None values
        thermal_params = [(k, v) for k, v in thermal_params if v is not None and v != ""]
        
        if thermal_params:
            if y < 150:
                c.showPage()
                y = height - 50
            
            # Section title
            c.setFont(font_bold, 12)
            params_title = "پارامترهای حرارتی" if is_rtl else "Thermal Parameters"
            if is_rtl:
                params_title = self._reshape_persian_text(params_title)
                self._draw_text_right_aligned(c, params_title, width - 50, y, is_rtl=False)
            else:
                c.drawString(50, y, params_title)
            y -= 20
            
            c.setFont(font_name, 10)
            col_switch = 0
            current_y = y
            
            for key, value in thermal_params:
                label = info_labels.get(key, key)
                
                # Format value with units
                if key == "emissivity":
                    display_value = f"{value}"
                elif key in ["ambient_temp", "reflected_temp"]:
                    display_value = f"{value}°C"
                elif key == "humidity":
                    display_value = f"{value}%"
                elif key == "distance":
                    display_value = f"{value}m"
                else:
                    display_value = str(value)
                
                if is_rtl:
                    text = f"{label}: {display_value}"
                    reshaped_text = self._reshape_persian_text(text)
                    
                    if col_switch % 2 == 0:
                        self._draw_text_right_aligned(c, reshaped_text, width / 2 - 20, current_y, is_rtl=False)
                    else:
                        self._draw_text_right_aligned(c, reshaped_text, width - 50, current_y, is_rtl=False)
                        current_y -= 15
                else:
                    if col_switch % 2 == 0:
                        c.drawString(col1_x, current_y, f"{label}: {display_value}")
                    else:
                        c.drawString(col2_x, current_y, f"{label}: {display_value}")
                        current_y -= 15
                
                col_switch += 1
            
            if col_switch % 2 == 1:
                current_y -= 15
            
            y = current_y - 10
        
        y = y - 10
        
        # Task/Description section (if provided)
        task_desc = metadata.get("task", metadata.get("description", ""))
        if task_desc:
            if y < 100:
                c.showPage()
                y = height - 50
            
            c.setFont(font_bold, 11)
            task_title = "وظیفه / شرح" if is_rtl else "Task / Description"
            if is_rtl:
                task_title = self._reshape_persian_text(task_title)
                self._draw_text_right_aligned(c, task_title, width - 50, y, is_rtl=False)
            else:
                c.drawString(50, y, task_title)
            y -= 18
            
            c.setFont(font_name, 9)
            # Wrap task text
            max_width = width - 100
            words = task_desc.split() if not is_rtl else self._reshape_persian_text(task_desc).split()
            line = ""
            
            for word in words:
                test_line = line + word + " "
                if c.stringWidth(test_line, font_name, 9) < max_width:
                    line = test_line
                else:
                    if line:
                        if is_rtl:
                            self._draw_text_right_aligned(c, line.strip(), width - 50, y, is_rtl=False)
                        else:
                            c.drawString(50, y, line.strip())
                        y -= 12
                    line = word + " "
            
            if line:
                if is_rtl:
                    self._draw_text_right_aligned(c, line.strip(), width - 50, y, is_rtl=False)
                else:
                    c.drawString(50, y, line.strip())
                y -= 15

        y -= 20
        
        print(f"[REPORT] Processing {len(images)} images")

        # Process images - show images and tables for markers and regions
        for idx, img_data in enumerate(images):
            print(f"[REPORT] Processing image {idx + 1}/{len(images)}: {img_data.get('name', 'unnamed')}")
            
            if y < 500:
                c.showPage()
                page_num += 1
                self._draw_professional_header(c, width, height, title, is_rtl, page_num)
                y = height - 100
            
            # Image Section Header
            img_title = f"تصویر {idx + 1}: {img_data.get('name', 'بدون نام')}" if is_rtl else f"Image {idx + 1}: {img_data.get('name', 'Unnamed')}"
            y = self._draw_section_header(c, img_title, y, width, is_rtl)

            # Load thermal and real images
            thermal_image = None
            real_image = None
            
            # Thermal image - support both path and base64
            if img_data.get("thermal_base64") or img_data.get("thermalBase64"):
                base64_data = img_data.get("thermal_base64") or img_data.get("thermalBase64")
                thermal_image = self._base64_to_image(base64_data)
                if thermal_image:
                    print(f"[REPORT] Loaded thermal image from base64")
            elif img_data.get("thermal_path") and os.path.exists(img_data["thermal_path"]):
                thermal_image = Image.open(img_data["thermal_path"])
                print(f"[REPORT] Loaded thermal image from path")

            # Real image - support both path and base64
            if img_data.get("real_base64") or img_data.get("realBase64"):
                base64_data = img_data.get("real_base64") or img_data.get("realBase64")
                real_image = self._base64_to_image(base64_data)
                if real_image:
                    print(f"[REPORT] Loaded real image from base64")
            elif img_data.get("real_path") and os.path.exists(img_data["real_path"]):
                real_image = Image.open(img_data["real_path"])
                print(f"[REPORT] Loaded real image from path")

            # Draw images side by side with borders
            img_width = 245
            img_height = 185
            spacing = 15
            border_padding = 5
            
            # Thermal image (left/right depending on RTL)
            if thermal_image:
                try:
                    img_reader = ImageReader(thermal_image)
                    thermal_x = 50 if not is_rtl else width - 50 - img_width - spacing - img_width
                    
                    # Border background
                    c.setFillColorRGB(0.95, 0.95, 0.95)
                    c.rect(thermal_x - border_padding, y - img_height - border_padding, 
                          img_width + 2*border_padding, img_height + 2*border_padding, fill=1, stroke=1)
                    c.setFillColorRGB(0, 0, 0)
                    
                    c.drawImage(img_reader, thermal_x, y - img_height, width=img_width, height=img_height)
                    
                    # Label below with background
                    c.setFont(font_bold if is_rtl and self.has_persian_font else 'Helvetica-Bold', 10)
                    label_text = "تصویر حرارتی" if is_rtl else "Thermal Image"
                    if is_rtl:
                        label_text = self._reshape_persian_text(label_text)
                    text_width = c.stringWidth(label_text, c._fontname, c._fontsize)
                    
                    # Background for label
                    c.setFillColorRGB(0.2, 0.5, 0.8)
                    c.rect(thermal_x, y - img_height - 25, img_width, 18, fill=1, stroke=0)
                    c.setFillColorRGB(1, 1, 1)
                    c.drawString(thermal_x + img_width/2 - text_width/2, y - img_height - 20, label_text)
                    c.setFillColorRGB(0, 0, 0)
                    
                    print(f"[REPORT] Thermal image added successfully")
                except Exception as e:
                    print(f"[REPORT] Error adding thermal image: {e}")

            # Real image (right/left depending on RTL)
            if real_image:
                try:
                    img_reader = ImageReader(real_image)
                    real_x = 50 + img_width + spacing if not is_rtl else width - 50 - img_width
                    
                    # Border background
                    c.setFillColorRGB(0.95, 0.95, 0.95)
                    c.rect(real_x - border_padding, y - img_height - border_padding, 
                          img_width + 2*border_padding, img_height + 2*border_padding, fill=1, stroke=1)
                    c.setFillColorRGB(0, 0, 0)
                    
                    c.drawImage(img_reader, real_x, y - img_height, width=img_width, height=img_height)
                    
                    # Label below with background
                    c.setFont(font_bold if is_rtl and self.has_persian_font else 'Helvetica-Bold', 10)
                    label_text = "تصویر واقعی" if is_rtl else "Real Image"
                    if is_rtl:
                        label_text = self._reshape_persian_text(label_text)
                    text_width = c.stringWidth(label_text, c._fontname, c._fontsize)
                    
                    # Background for label
                    c.setFillColorRGB(0.2, 0.5, 0.8)
                    c.rect(real_x, y - img_height - 25, img_width, 18, fill=1, stroke=0)
                    c.setFillColorRGB(1, 1, 1)
                    c.drawString(real_x + img_width/2 - text_width/2, y - img_height - 20, label_text)
                    c.setFillColorRGB(0, 0, 0)
                    
                    print(f"[REPORT] Real image added successfully")
                except Exception as e:
                    print(f"[REPORT] Error adding real image: {e}")
            
            y -= img_height + 35
            
            # Get markers for this image
            img_markers = [m for m in markers if m.get('image_id') == img_data.get('id')]
            
            # Get regions for this image  
            img_regions = [r for r in regions if r.get('image_id') == img_data.get('id')]
            
            # Markers table
            if img_markers:
                if y < 200:
                    c.showPage()
                    page_num += 1
                    self._draw_professional_header(c, width, height, title, is_rtl, page_num)
                    y = height - 100
                
                # Section header
                y = self._draw_section_header(c, 
                    "نشانگرهای دما" if is_rtl else "Temperature Markers", 
                    y, width, is_rtl)
                
                # Create table data
                if is_rtl:
                    table_data = [["دما (°C)", "Y", "X", "نام"]]
                    for marker in img_markers:
                        table_data.append([
                            f"{marker.get('temperature', 0):.2f}",
                            f"{marker.get('y', 0):.1f}",
                            f"{marker.get('x', 0):.1f}",
                            marker.get('label', '')
                        ])
                    col_widths = [90, 70, 70, 160]
                else:
                    table_data = [["Label", "X", "Y", "Temp (°C)"]]
                    for marker in img_markers:
                        table_data.append([
                            marker.get('label', ''),
                            f"{marker.get('x', 0):.1f}",
                            f"{marker.get('y', 0):.1f}",
                            f"{marker.get('temperature', 0):.2f}"
                        ])
                    col_widths = [160, 70, 70, 90]
                
                # Draw enhanced table
                y = self._draw_enhanced_table(c, table_data, 50, y, col_widths, is_rtl)
                y -= 20
            
            # Regions table
            if img_regions:
                if y < 200:
                    c.showPage()
                    page_num += 1
                    self._draw_professional_header(c, width, height, title, is_rtl, page_num)
                    y = height - 100
                
                # Section header
                y = self._draw_section_header(c, 
                    "مناطق تحلیل" if is_rtl else "Analysis Regions", 
                    y, width, is_rtl)
                
                # Create regions table
                if is_rtl:
                    region_data = [["میانگین (°C)", "حداکثر (°C)", "حداقل (°C)", "نوع", "نام"]]
                    for region in img_regions:
                        region_data.append([
                            f"{region.get('avg_temp', 0):.2f}",
                            f"{region.get('max_temp', 0):.2f}",
                            f"{region.get('min_temp', 0):.2f}",
                            region.get('type', ''),
                            region.get('label', '')
                        ])
                    col_widths = [90, 90, 90, 70, 130]
                else:
                    region_data = [["Label", "Type", "Min (°C)", "Max (°C)", "Avg (°C)"]]
                    for region in img_regions:
                        region_data.append([
                            region.get('label', ''),
                            region.get('type', ''),
                            f"{region.get('min_temp', 0):.2f}",
                            f"{region.get('max_temp', 0):.2f}",
                            f"{region.get('avg_temp', 0):.2f}"
                        ])
                    col_widths = [130, 70, 90, 90, 90]
                
                # Draw enhanced regions table
                y = self._draw_enhanced_table(c, region_data, 50, y, col_widths, is_rtl)
                y -= 25
        
        # Temperature Histogram (overall - from CSV if available, otherwise from markers)
        histogram_buffer = None
        
        # Try to create histogram from CSV data first (more accurate)
        for img_data in images:
            csv_url = img_data.get('csv_url') or img_data.get('csvUrl')
            if csv_url:
                try:
                    # Read CSV file
                    import requests
                    from pathlib import Path
                    
                    # Check if it's a local file path
                    if csv_url.startswith('/files/'):
                        # Local file - construct path
                        csv_path = Path(csv_url.replace('/files/', str(Path(__file__).resolve().parents[3]) + '/'))
                        if csv_path.exists():
                            with open(csv_path, 'r') as f:
                                csv_data = f.read()
                            histogram_buffer = self._create_temperature_histogram_from_csv(csv_data, language)
                            if histogram_buffer:
                                print(f"[REPORT] Histogram created from CSV: {csv_url}")
                                break
                except Exception as e:
                    print(f"[REPORT] Could not create histogram from CSV: {e}")
        
        # Fallback to markers if CSV histogram failed
        if not histogram_buffer and markers:
            print(f"[REPORT] Creating histogram from markers")
            histogram_buffer = self._create_temperature_histogram(markers, language)
        
        if histogram_buffer:
            print(f"[REPORT] Adding temperature histogram")
            
            if y < 300:
                c.showPage()
                page_num += 1
                self._draw_professional_header(c, width, height, title, is_rtl, page_num)
                y = height - 100
            
            # Section header
            y = self._draw_section_header(c, 
                "توزیع دما" if is_rtl else "Temperature Distribution", 
                y, width, is_rtl)
            
            try:
                histogram_img = Image.open(histogram_buffer)
                img_reader = ImageReader(histogram_img)
                
                # Draw histogram centered with border
                hist_width = 450
                hist_height = 220
                hist_x = (width - hist_width) / 2
                
                # Background box
                c.setFillColorRGB(0.98, 0.98, 0.98)
                c.rect(hist_x - 10, y - hist_height - 10, hist_width + 20, hist_height + 20, fill=1, stroke=1)
                c.setFillColorRGB(0, 0, 0)
                
                c.drawImage(img_reader, hist_x, y - hist_height, width=hist_width, height=hist_height)
                y -= hist_height + 40
                print(f"[REPORT] Histogram added successfully")
            except Exception as e:
                print(f"[REPORT] Error adding histogram: {e}")
        
        # Notes Section
        if notes:
            print(f"[REPORT] Adding notes")
            
            if y < 200:
                c.showPage()
                page_num += 1
                self._draw_professional_header(c, width, height, title, is_rtl, page_num)
                y = height - 100
            
            # Section header
            y = self._draw_section_header(c, 
                "یادداشت‌ها" if is_rtl else "Notes", 
                y, width, is_rtl)
            
            c.setFont(font_name, 10)
            
            if is_rtl:
                # For RTL text, reshape the entire notes first
                reshaped_notes = self._reshape_persian_text(notes)
                # Simple text wrapping for RTL
                max_width = 500
                words = reshaped_notes.split()
                line = ""
                for word in words:
                    test_line = line + word + " "
                    if c.stringWidth(test_line, font_name, 10) < max_width:
                        line = test_line
                    else:
                        if line:
                            self._draw_text_right_aligned(c, line.strip(), width - 50, y, is_rtl=False)
                            y -= 15
                            if y < 100:
                                c.showPage()
                                y = height - 50
                        line = word + " "
                if line:
                    self._draw_text_right_aligned(c, line.strip(), width - 50, y, is_rtl=False)
            else:
                # LTR text wrapping
                max_width = 500
                words = notes.split()
                line = ""
                for word in words:
                    test_line = line + word + " "
                    if c.stringWidth(test_line, font_name, 10) < max_width:
                        line = test_line
                    else:
                        if line:
                            c.drawString(50, y, line.strip())
                            y -= 15
                            if y < 100:
                                c.showPage()
                                y = height - 50
                        line = word + " "
                if line:
                    c.drawString(50, y, line.strip())
        
        # Professional Footer on last page
        date_str = self._get_formatted_date(language)
        self._draw_professional_footer(c, width, date_str, is_rtl)
        
        # Save PDF
        c.save()
        
        print(f"[REPORT] PDF generated successfully: {pdf_path}")
        return str(pdf_path)

    def generate_docx_report(
        self,
        project_id: str,
        metadata: Dict[str, Any],
        images: List[Dict[str, Any]],
        markers: List[Dict[str, Any]],
        regions: List[Dict[str, Any]],
        notes: Optional[str] = None,
        language: str = "en"
    ) -> str:
        """
        Generate DOCX report with bilingual support (English/Persian)
        Returns: Path to generated DOCX file
        """
        print(f"[REPORT] Starting DOCX generation for project {project_id} in {language}")
        
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Create report directory
        dirs = self.file_manager.create_project_directories(project_id)
        if not dirs:
            raise Exception("Failed to create project directories")
            
        report_dir = dirs["reports"]
        print(f"[REPORT] Report directory: {report_dir}")

        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        docx_filename = f"report_{timestamp}.docx"
        docx_path = report_dir / docx_filename

        # Create document
        document = Document()
        
        # Set document properties
        core_properties = document.core_properties
        core_properties.title = metadata.get("title", "Thermal Analysis Report")
        core_properties.author = "Thermal Analyzer Pro"

        # Professional Title with background
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(metadata.get("title", "Thermal Analysis Report"))
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(26, 77, 153)  # Dark blue
        
        # Subtitle
        subtitle_para = document.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("Thermal Analysis Report")
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(50, 127, 204)
        
        # Separator line
        document.add_paragraph("_" * 80)

        # Project Information Section with styled heading
        section_heading = document.add_heading("Project Information", level=1)
        section_heading.runs[0].font.color.rgb = RGBColor(26, 77, 153)
        
        info_table = document.add_table(rows=8, cols=2)
        info_table.style = 'Medium Shading 1 Accent 1'

        info_data = [
            ("Project Name", metadata.get("project_name", "")),
            ("Company", metadata.get("company", "")),
            ("Operator", metadata.get("operator", "")),
            ("Customer", metadata.get("customer", "")),
            ("Device", metadata.get("device", metadata.get("cameraModel", ""))),
            ("Serial Number", metadata.get("serial_no", metadata.get("serial_number", ""))),
            ("Location", metadata.get("measuring_site", metadata.get("location", ""))),
            ("Date", metadata.get("date", self._get_formatted_date(language))),
        ]

        for idx, (label, value) in enumerate(info_data):
            row_cells = info_table.rows[idx].cells
            # Label cell - bold
            label_run = row_cells[0].paragraphs[0].add_run(label)
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            # Value cell
            value_run = row_cells[1].paragraphs[0].add_run(str(value))
            value_run.font.size = Pt(11)

        # Images Section
        if images:
            print(f"[REPORT] Adding {len(images)} images to DOCX")
            document.add_page_break()
            
            # Section heading
            img_section_heading = document.add_heading("Thermal Images", level=1)
            img_section_heading.runs[0].font.color.rgb = RGBColor(26, 77, 153)

            for idx, img_data in enumerate(images):
                print(f"[REPORT] Processing image {idx + 1}/{len(images)}: {img_data.get('name', 'unnamed')}")
                
                # Image subsection
                img_heading = document.add_heading(f"Image {idx + 1}: {img_data.get('name', 'Unnamed')}", level=2)
                img_heading.runs[0].font.color.rgb = RGBColor(50, 127, 204)

                # Create table for side-by-side images
                img_table = document.add_table(rows=2, cols=2)
                img_table.style = 'Light Grid'
                
                # Labels
                img_table.rows[0].cells[0].text = "Thermal Image"
                img_table.rows[0].cells[1].text = "Real Image"
                for cell in img_table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(26, 77, 153)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Thermal image
                thermal_image = None
                if img_data.get("thermal_base64") or img_data.get("thermalBase64"):
                    base64_data = img_data.get("thermal_base64") or img_data.get("thermalBase64")
                    thermal_image = self._base64_to_image(base64_data)
                    if thermal_image:
                        print(f"[REPORT] Loaded thermal image from base64")
                elif img_data.get("thermal_path") and os.path.exists(img_data["thermal_path"]):
                    thermal_image = Image.open(img_data["thermal_path"])
                    print(f"[REPORT] Loaded thermal image from path")

                if thermal_image:
                    try:
                        temp_thermal_path = report_dir / f"temp_thermal_{img_data.get('id', idx)}.png"
                        thermal_image.save(temp_thermal_path)
                        # Add to table cell
                        para = img_table.rows[1].cells[0].paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.add_run().add_picture(str(temp_thermal_path), width=Inches(2.8))
                        temp_thermal_path.unlink()
                        print(f"[REPORT] Thermal image added successfully")
                    except Exception as e:
                        print(f"[REPORT] Error adding thermal image to DOCX: {e}")
                        img_table.rows[1].cells[0].text = f"[Image error: {e}]"

                # Real image
                real_image = None
                if img_data.get("real_base64") or img_data.get("realBase64"):
                    base64_data = img_data.get("real_base64") or img_data.get("realBase64")
                    real_image = self._base64_to_image(base64_data)
                    if real_image:
                        print(f"[REPORT] Loaded real image from base64")
                elif img_data.get("real_path") and os.path.exists(img_data["real_path"]):
                    real_image = Image.open(img_data["real_path"])
                    print(f"[REPORT] Loaded real image from path")

                if real_image:
                    try:
                        temp_real_path = report_dir / f"temp_real_{img_data.get('id', idx)}.png"
                        real_image.save(temp_real_path)
                        # Add to table cell
                        para = img_table.rows[1].cells[1].paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.add_run().add_picture(str(temp_real_path), width=Inches(2.8))
                        temp_real_path.unlink()
                        print(f"[REPORT] Real image added successfully")
                    except Exception as e:
                        print(f"[REPORT] Error adding real image to DOCX: {e}")
                        img_table.rows[1].cells[1].text = f"[Image error: {e}]"

                document.add_paragraph()  # Add spacing

        # Markers Section
        if markers:
            print(f"[REPORT] Adding {len(markers)} markers to DOCX")
            document.add_page_break()
            
            # Section heading
            markers_heading = document.add_heading("Temperature Markers", level=1)
            markers_heading.runs[0].font.color.rgb = RGBColor(26, 77, 153)

            markers_table = document.add_table(rows=len(markers) + 1, cols=4)
            markers_table.style = 'Medium Shading 1 Accent 1'

            # Headers with styling
            headers = ["Label", "Position (X, Y)", "Temperature (°C)", "Image ID"]
            for idx, header in enumerate(headers):
                cell = markers_table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            # Data
            for idx, marker in enumerate(markers, 1):
                try:
                    markers_table.rows[idx].cells[0].text = marker.get('label', '')
                    markers_table.rows[idx].cells[1].text = f"({marker.get('x', 0):.1f}, {marker.get('y', 0):.1f})"
                    
                    # Temperature cell - color code if extreme
                    temp = marker.get('temperature', 0)
                    temp_cell = markers_table.rows[idx].cells[2]
                    temp_cell.text = f"{temp:.2f}"
                    
                    markers_table.rows[idx].cells[3].text = marker.get('image_id', '')[:8]
                except Exception as e:
                    print(f"[REPORT] Error adding marker {idx}: {e}")

        # Regions Section
        if regions:
            print(f"[REPORT] Adding {len(regions)} regions to DOCX")
            document.add_page_break()
            
            # Section heading
            regions_heading = document.add_heading("Analysis Regions", level=1)
            regions_heading.runs[0].font.color.rgb = RGBColor(26, 77, 153)

            regions_table = document.add_table(rows=len(regions) + 1, cols=5)
            regions_table.style = 'Medium Shading 1 Accent 1'

            # Headers with styling
            headers = ["Label", "Type", "Min Temp (°C)", "Max Temp (°C)", "Avg Temp (°C)"]
            for idx, header in enumerate(headers):
                cell = regions_table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

            # Data
            for idx, region in enumerate(regions, 1):
                try:
                    regions_table.rows[idx].cells[0].text = region.get('label', '')
                    regions_table.rows[idx].cells[1].text = region.get('type', '')
                    regions_table.rows[idx].cells[2].text = f"{region.get('min_temp', 0):.2f}"
                    regions_table.rows[idx].cells[3].text = f"{region.get('max_temp', 0):.2f}"
                    regions_table.rows[idx].cells[4].text = f"{region.get('avg_temp', 0):.2f}"
                except Exception as e:
                    print(f"[REPORT] Error adding region {idx}: {e}")

        # Notes Section
        if notes:
            print(f"[REPORT] Adding notes to DOCX")
            document.add_page_break()
            
            # Section heading
            notes_heading = document.add_heading("Notes", level=1)
            notes_heading.runs[0].font.color.rgb = RGBColor(26, 77, 153)
            
            notes_para = document.add_paragraph(notes)
            notes_para.runs[0].font.size = Pt(11)

        # Professional Footer
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Footer text with styling
        footer_run = footer_para.add_run(f"Generated by Thermal Analyzer Pro - {self._get_formatted_date(language)} {datetime.utcnow().strftime('%H:%M:%S')}")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(100, 100, 100)

        # Save document
        document.save(str(docx_path))
        
        print(f"[REPORT] DOCX generated successfully: {docx_path}")
        return str(docx_path)